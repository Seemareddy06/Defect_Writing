import streamlit as st
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt

# ---------------------------
# CONFIGURATION
# ---------------------------
st.set_page_config(page_title="AI Jira Defect Writer", layout="wide")
st.title("🧠 AI-Powered Jira Defect Writing Tool (Enhanced)")
st.markdown("Paste your user story, issue, and impact area to generate a detailed, Jira-ready defect report automatically.")

# ---------------------------
# SIDEBAR INPUTS
# ---------------------------
OPENROUTER_API_KEY = "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"  # replace with your valid key

with st.sidebar:
    st.header("Defect Details")
    sprint_number = st.number_input("🏷️ Sprint Number", min_value=1, step=1)
    module_name = st.text_input("Module Name", placeholder="e.g., ICHRA")
    environment = st.selectbox("🌐 Environment", ["QA", "UAT", "Production"])
    group_id = st.text_input("Group ID", placeholder="Enter Group ID")
    plan_id = st.text_input("Plan ID", placeholder="Enter Plan ID")

# ---------------------------
# MAIN BODY INPUT
# ---------------------------
user_story = st.text_area(
    "Paste User Story / Issue Details",
    placeholder="Describe the issue, example: 'Rate calculation is incorrect when selecting age banded plan with dependent slider missing...'",
    height=200
)

impact_area = st.text_area(
    "Impact Area / Additional Context",
    placeholder="Describe which part of the system is impacted and any relevant details (e.g., Enrollment workflow, Rate calculation, UI, etc.)",
    height=100
)

generate_btn = st.button("🚀 Generate Jira Defect")

# ---------------------------
# AI DEFECT GENERATION
# ---------------------------
if generate_btn:
    if not user_story.strip() or not module_name.strip():
        st.warning("⚠️ Please enter both Module Name and User Story / Issue details.")
    else:
        # Dynamically include impact area line only if provided
        impact_sentence = f"This impacts the {impact_area.strip()}." if impact_area.strip() else ""

        # ---------------------------
        # PROMPT TEMPLATE
        # ---------------------------
        system_prompt = f"""
        You are a senior QA Engineer writing Jira defects in a structured, professional format.
        Based on the provided user story, issue details, and impact area, generate a **complete Jira defect**.

        Follow this EXACT format:
        TITLE: Sprint {int(sprint_number)} - {module_name} - <short issue summary>

        ISSUE DESCRIPTION:
        As a Superuser or Broker user in {environment} environment, {module_name} module has the following issue:
        <Describe what’s happening, with context and technical impact.>
        {impact_sentence}

        NAVIGATION PATH:
        Provide the exact path or sequence of screens/menus to reach the defect, e.g.,
        Superuser → Group → Add Group → {module_name} → Plan Selection → Save and Next.

        STEPS TO REPRODUCE:
        1. Login as Superuser/Broker user.
        2. Navigate to the relevant module path.
        3. Perform the necessary actions mentioned in the user story.
        4. Observe the incorrect or unexpected behavior.

        EXPECTED RESULT:
        Describe what the system should ideally do (correct calculation, proper display, successful save, etc.).

        ACTUAL RESULT:
        Describe what actually happens (wrong calculation, missing slider, error message, etc.).

        PLAN ID: {plan_id}
        GROUP ID: {group_id}

        Notes:
        - Keep all sentences clear, concise, and professional.
        - Avoid filler words.
        - Automatically infer the issue type (UI, calculation, validation, dependency, etc.).
        - Ensure the language matches Jira defect standards.
        """

        user_prompt = f"""
        Module: {module_name}
        Environment: {environment}
        User Story / Issue Context:
        {user_story}

        Impact Area:
        {impact_area}
        """

        # ---------------------------
        # API CALL
        # ---------------------------
        with st.spinner("🤖 Generating Jira defect report..."):
            try:
                url = "https://openrouter.ai/api/v1/chat/completions"
                headers = {
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json",
                }
                payload = {
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "temperature": 0.3,
                }

                response = requests.post(url, headers=headers, json=payload)
                response.raise_for_status()
                data = response.json()

                message = data.get("choices", [{}])[0].get("message", {}).get("content", "")
                ai_output = message.strip()

                if not ai_output:
                    st.error("❌ AI returned empty output. Please refine your input.")
                else:
                    st.success("✅ Jira Defect Generated Successfully!")

                    # --- Display ---
                    st.text_area("📄 Jira Defect Output", ai_output, height=400)

                    # --- Export to Word ---
                    doc = Document()
                    doc.add_heading(f"Defect Report (Sprint {int(sprint_number)})", level=1)
                    for line in ai_output.splitlines():
                        if line.strip().startswith(("TITLE:", "ISSUE DESCRIPTION:", "NAVIGATION PATH:",
                                                    "STEPS TO REPRODUCE:", "EXPECTED RESULT:",
                                                    "ACTUAL RESULT:", "PLAN ID:", "GROUP ID:")):
                            p = doc.add_paragraph(line.strip())
                            p.runs[0].bold = True
                            p.runs[0].font.size = Pt(12)
                        elif line.strip():
                            doc.add_paragraph(line.strip())

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        label="⬇️ Download Jira Defect (Word)",
                        data=buffer,
                        file_name=f"sprint_{int(sprint_number)}_jira_defect.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            except Exception as e:
                st.error(f"❌ Error generating report: {e}")
